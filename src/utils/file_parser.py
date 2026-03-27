#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文件解析工具
支持解析Word、Excel、PDF、CSV、JPG等多种文件格式
"""

import os
import csv
import json
import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from io import BytesIO

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from PIL import Image
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False


class FileParser:
    """文件解析器基类"""

    def __init__(self, file_path: str = None, file_content: bytes = None, filename: str = None):
        """
        初始化文件解析器

        Args:
            file_path: 文件路径
            file_content: 文件内容（bytes）
            filename: 文件名
        """
        self.file_path = file_path
        self.file_content = file_content
        self.filename = filename
        self.file_ext = self._get_file_ext()

    def _get_file_ext(self) -> str:
        """获取文件扩展名"""
        if self.filename:
            return Path(self.filename).suffix.lower()
        if self.file_path:
            return Path(self.file_path).suffix.lower()
        return ''

    def parse(self) -> Dict[str, Any]:
        """
        解析文件

        Returns:
            解析后的数据字典
        """
        if not self.file_ext:
            raise ValueError("无法确定文件类型")

        if self.file_ext == '.csv':
            return self.parse_csv()
        elif self.file_ext in ['.xls', '.xlsx']:
            return self.parse_excel()
        elif self.file_ext == '.pdf':
            return self.parse_pdf()
        elif self.file_ext in ['.doc', '.docx']:
            return self.parse_word()
        elif self.file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return self.parse_image()
        elif self.file_ext == '.json':
            return self.parse_json()
        elif self.file_ext == '.txt':
            return self.parse_txt()
        else:
            raise ValueError(f"不支持的文件类型: {self.file_ext}")

    def parse_csv(self) -> Dict[str, Any]:
        """解析CSV文件"""
        try:
            if self.file_content:
                content = self.file_content.decode('utf-8-sig')
                csv_reader = csv.reader(content.splitlines())
            else:
                with open(self.file_path, 'r', encoding='utf-8-sig') as f:
                    csv_reader = csv.reader(f)

            data = list(csv_reader)
            headers = data[0] if data else []
            rows = data[1:] if len(data) > 1 else []

            return {
                'type': 'csv',
                'headers': headers,
                'rows': rows,
                'total_rows': len(rows),
                'success': True
            }
        except Exception as e:
            return {
                'type': 'csv',
                'success': False,
                'error': str(e)
            }

    def parse_excel(self) -> Dict[str, Any]:
        """解析Excel文件"""
        if not EXCEL_AVAILABLE:
            return {
                'type': 'excel',
                'success': False,
                'error': 'openpyxl未安装，请先安装: pip install openpyxl'
            }

        try:
            if self.file_content:
                workbook = openpyxl.load_workbook(BytesIO(self.file_content))
            else:
                workbook = openpyxl.load_workbook(self.file_path)

            sheets_data = {}

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                data = []

                for row in sheet.iter_rows(values_only=True):
                    data.append([cell if cell is not None else '' for cell in row])

                if data:
                    headers = data[0]
                    rows = data[1:] if len(data) > 1 else []
                    sheets_data[sheet_name] = {
                        'headers': headers,
                        'rows': rows,
                        'total_rows': len(rows)
                    }

            return {
                'type': 'excel',
                'sheets': sheets_data,
                'total_sheets': len(sheets_data),
                'success': True
            }
        except Exception as e:
            return {
                'type': 'excel',
                'success': False,
                'error': str(e)
            }

    def parse_pdf(self) -> Dict[str, Any]:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            return {
                'type': 'pdf',
                'success': False,
                'error': 'PyMuPDF未安装，请先安装: pip install pymupdf'
            }

        try:
            if self.file_content:
                doc = fitz.open(stream=BytesIO(self.file_content), filetype="pdf")
            else:
                doc = fitz.open(self.file_path)

            text = ""
            pages = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                text += page_text + "\n"
                pages.append({
                    'page_num': page_num + 1,
                    'text': page_text
                })

            doc.close()

            return {
                'type': 'pdf',
                'text': text.strip(),
                'pages': pages,
                'total_pages': len(pages),
                'success': True
            }
        except Exception as e:
            return {
                'type': 'pdf',
                'success': False,
                'error': str(e)
            }

    def parse_word(self) -> Dict[str, Any]:
        """解析Word文件"""
        if not WORD_AVAILABLE:
            return {
                'type': 'word',
                'success': False,
                'error': 'python-docx未安装，请先安装: pip install python-docx'
            }

        try:
            if self.file_content:
                doc = Document(BytesIO(self.file_content))
            else:
                doc = Document(self.file_path)

            paragraphs = []
            tables = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            text = '\n'.join(paragraphs)

            return {
                'type': 'word',
                'text': text,
                'paragraphs': paragraphs,
                'tables': tables,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'success': True
            }
        except Exception as e:
            return {
                'type': 'word',
                'success': False,
                'error': str(e)
            }

    def parse_image(self) -> Dict[str, Any]:
        """解析图片文件"""
        if not IMAGE_AVAILABLE:
            return {
                'type': 'image',
                'success': False,
                'error': 'Pillow未安装，请先安装: pip install Pillow'
            }

        try:
            if self.file_content:
                img = Image.open(BytesIO(self.file_content))
            else:
                img = Image.open(self.file_path)

            info = {
                'format': img.format,
                'mode': img.mode,
                'size': img.size,
                'width': img.width,
                'height': img.height
            }

            # 如果图片包含文字（OCR），这里可以添加OCR功能
            # 需要安装: pip install pytesseract

            return {
                'type': 'image',
                'info': info,
                'success': True
            }
        except Exception as e:
            return {
                'type': 'image',
                'success': False,
                'error': str(e)
            }

    def parse_json(self) -> Dict[str, Any]:
        """解析JSON文件"""
        try:
            if self.file_content:
                data = json.loads(self.file_content.decode('utf-8'))
            else:
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)

            return {
                'type': 'json',
                'data': data,
                'success': True
            }
        except Exception as e:
            return {
                'type': 'json',
                'success': False,
                'error': str(e)
            }

    def parse_txt(self) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            if self.file_content:
                text = self.file_content.decode('utf-8')
            else:
                with open(self.file_path, 'r', encoding='utf-8') as f:
                    text = f.read()

            return {
                'type': 'txt',
                'text': text,
                'success': True
            }
        except Exception as e:
            return {
                'type': 'txt',
                'success': False,
                'error': str(e)
            }


def parse_file(file_path: str = None, file_content: bytes = None, filename: str = None) -> Dict[str, Any]:
    """
    解析文件的便捷函数

    Args:
        file_path: 文件路径
        file_content: 文件内容（bytes）
        filename: 文件名

    Returns:
        解析后的数据字典
    """
    parser = FileParser(file_path, file_content, filename)
    return parser.parse()


def extract_product_data_from_file(file_path: str = None, file_content: bytes = None, 
                                   filename: str = None) -> Dict[str, Any]:
    """
    从文件中提取产品数据

    Args:
        file_path: 文件路径
        file_content: 文件内容（bytes）
        filename: 文件名

    Returns:
        提取的产品数据
    """
    parsed_data = parse_file(file_path, file_content, filename)

    if not parsed_data.get('success'):
        return {
            'success': False,
            'error': parsed_data.get('error', '解析失败')
        }

    file_type = parsed_data.get('type')

    # 根据文件类型提取产品数据
    if file_type == 'csv':
        return extract_products_from_csv(parsed_data)
    elif file_type == 'excel':
        return extract_products_from_excel(parsed_data)
    elif file_type == 'pdf':
        return extract_products_from_text(parsed_data.get('text', ''))
    elif file_type == 'word':
        return extract_products_from_text(parsed_data.get('text', ''))
    elif file_type == 'json':
        return extract_products_from_json(parsed_data.get('data', {}))
    else:
        return {
            'success': False,
            'error': f'不支持从 {file_type} 文件类型提取产品数据'
        }


def extract_products_from_csv(parsed_data: Dict) -> Dict[str, Any]:
    """从CSV数据中提取产品"""
    headers = parsed_data.get('headers', [])
    rows = parsed_data.get('rows', [])

    products = []

    # 尝试识别列名
    airline_col = None
    product_name_col = None
    route_col = None
    booking_class_col = None
    price_col = None
    rebate_col = None
    office_col = None
    remarks_col = None
    valid_period_col = None

    for header in headers:
        header_lower = header.lower()
        if 'airline' in header_lower or '航司' in header_lower or '航空公司' in header_lower:
            airline_col = headers.index(header)
        elif 'product' in header_lower or '产品' in header_lower:
            product_name_col = headers.index(header)
        elif 'route' in header_lower or '航线' in header_lower:
            route_col = headers.index(header)
        elif 'class' in header_lower or '舱位' in header_lower:
            booking_class_col = headers.index(header)
        elif 'price' in header_lower or '价格' in header_lower:
            price_col = headers.index(header)
        elif 'rebate' in header_lower or '返点' in header_lower:
            rebate_col = headers.index(header)
        elif 'office' in header_lower:
            office_col = headers.index(header)
        elif 'remarks' in header_lower or '备注' in header_lower:
            remarks_col = headers.index(header)
        elif 'valid' in header_lower or '有效期' in header_lower or 'period' in header_lower:
            valid_period_col = headers.index(header)

    for row in rows:
        if len(row) < len(headers):
            continue

        product = {
            'airline': row[airline_col] if airline_col is not None else '',
            'product_name': row[product_name_col] if product_name_col is not None else '',
            'route': row[route_col] if route_col is not None else '',
            'booking_class': row[booking_class_col] if booking_class_col is not None else '',
            'price_increase': row[price_col] if price_col is not None else 0,
            'rebate_ratio': row[rebate_col] if rebate_col is not None else '',
            'office': row[office_col] if office_col is not None else '',
            'remarks': row[remarks_col] if remarks_col is not None else '',
            'valid_period': row[valid_period_col] if valid_period_col is not None else '',
            'ticket_type': 'ALL',
            'policy_identifier': '',
            'policy_code': ''
        }

        # 过滤空产品
        if product['airline'] and product['product_name']:
            products.append(product)

    return {
        'success': True,
        'type': 'csv',
        'products': products,
        'total_products': len(products)
    }


def extract_products_from_excel(parsed_data: Dict) -> Dict[str, Any]:
    """从Excel数据中提取产品"""
    sheets = parsed_data.get('sheets', {})
    all_products = []

    for sheet_name, sheet_data in sheets.items():
        headers = sheet_data.get('headers', [])
        rows = sheet_data.get('rows', [])

        # 尝试识别列名
        airline_col = None
        product_name_col = None
        route_col = None
        booking_class_col = None
        price_col = None
        rebate_col = None
        office_col = None
        remarks_col = None
        valid_period_col = None

        for header in headers:
            header_lower = str(header).lower()
            if 'airline' in header_lower or '航司' in header_lower or '航空公司' in header_lower:
                airline_col = headers.index(header)
            elif 'product' in header_lower or '产品' in header_lower:
                product_name_col = headers.index(header)
            elif 'route' in header_lower or '航线' in header_lower:
                route_col = headers.index(header)
            elif 'class' in header_lower or '舱位' in header_lower:
                booking_class_col = headers.index(header)
            elif 'price' in header_lower or '价格' in header_lower:
                price_col = headers.index(header)
            elif 'rebate' in header_lower or '返点' in header_lower:
                rebate_col = headers.index(header)
            elif 'office' in header_lower:
                office_col = headers.index(header)
            elif 'remarks' in header_lower or '备注' in header_lower:
                remarks_col = headers.index(header)
            elif 'valid' in header_lower or '有效期' in header_lower or 'period' in header_lower:
                valid_period_col = headers.index(header)

        for row in rows:
            if len(row) < len(headers):
                continue

            product = {
                'airline': str(row[airline_col]) if airline_col is not None and row[airline_col] else '',
                'product_name': str(row[product_name_col]) if product_name_col is not None and row[product_name_col] else '',
                'route': str(row[route_col]) if route_col is not None and row[route_col] else '',
                'booking_class': str(row[booking_class_col]) if booking_class_col is not None and row[booking_class_col] else '',
                'price_increase': row[price_col] if price_col is not None else 0,
                'rebate_ratio': str(row[rebate_col]) if rebate_col is not None and row[rebate_col] else '',
                'office': str(row[office_col]) if office_col is not None and row[office_col] else '',
                'remarks': str(row[remarks_col]) if remarks_col is not None and row[remarks_col] else '',
                'valid_period': str(row[valid_period_col]) if valid_period_col is not None and row[valid_period_col] else '',
                'ticket_type': 'ALL',
                'policy_identifier': '',
                'policy_code': ''
            }

            # 过滤空产品
            if product['airline'] and product['product_name']:
                all_products.append(product)

    return {
        'success': True,
        'type': 'excel',
        'products': all_products,
        'total_products': len(all_products)
    }


def extract_products_from_text(text: str) -> Dict[str, Any]:
    """从文本中提取产品数据"""
    # 这里可以实现从PDF、Word等文本文件中提取产品数据的逻辑
    # 使用正则表达式等工具来识别产品信息

    products = []

    # 简单的实现：查找包含航司代码的行
    lines = text.split('\n')
    airline_codes = ['CA', 'MU', 'CZ', 'HU', '3U', 'SC', 'MF', 'ZH', 'EU', '8L', 
                   'KY', 'PN', 'JR', 'GJ', '9H', 'UQ', 'GS', 'HO', 'JD', 'A6',
                   'DR', 'G5', 'BK', 'DZ', 'FU', 'GX', 'TV', 'RY', 'QW', 'NS']

    for line in lines:
        for code in airline_codes:
            if code in line:
                products.append({
                    'airline': code,
                    'product_name': line[:100],  # 取前100个字符作为产品名
                    'route': '',
                    'booking_class': '',
                    'price_increase': 0,
                    'rebate_ratio': '',
                    'office': '',
                    'remarks': line,
                    'valid_period': '',
                    'ticket_type': 'ALL',
                    'policy_identifier': '',
                    'policy_code': ''
                })
                break

    return {
        'success': True,
        'type': 'text',
        'products': products,
        'total_products': len(products),
        'warning': '从文本文件提取的产品数据可能不完整，建议使用CSV或Excel格式'
    }


def extract_products_from_json(data: dict) -> Dict[str, Any]:
    """从JSON数据中提取产品"""
    products = []

    # 尝试从JSON中提取产品列表
    if isinstance(data, dict):
        if 'products' in data:
            products = data['products']
        elif 'data' in data and isinstance(data['data'], list):
            products = data['data']
        elif isinstance(data, list):
            products = data
    elif isinstance(data, list):
        products = data

    # 标准化产品数据
    standardized_products = []
    for item in products:
        if isinstance(item, dict):
            product = {
                'airline': item.get('airline', ''),
                'product_name': item.get('product_name', item.get('name', '')),
                'route': item.get('route', item.get('航线', '')),
                'booking_class': item.get('booking_class', item.get('舱位', '')),
                'price_increase': item.get('price_increase', item.get('price', 0)),
                'rebate_ratio': item.get('rebate_ratio', item.get('rebate', '')),
                'office': item.get('office', ''),
                'remarks': item.get('remarks', item.get('备注', '')),
                'valid_period': item.get('valid_period', item.get('有效期', '')),
                'ticket_type': item.get('ticket_type', 'ALL'),
                'policy_identifier': item.get('policy_identifier', ''),
                'policy_code': item.get('policy_code', '')
            }

            if product['airline'] and product['product_name']:
                standardized_products.append(product)

    return {
        'success': True,
        'type': 'json',
        'products': standardized_products,
        'total_products': len(standardized_products)
    }
